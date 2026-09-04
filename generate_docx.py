import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()

    # Define Theme Colors
    COLOR_NAVY = RGBColor(31, 78, 120)     # #1F4E78 - Primary Heading
    COLOR_SECONDARY = RGBColor(46, 117, 182) # #2E75B6 - Secondary Heading
    COLOR_DARK = RGBColor(38, 38, 38)       # #262626 - Text
    COLOR_MUTED = RGBColor(89, 89, 89)     # #595959 - Subtitle / Muted

    HEX_PRIMARY = "1F4E78"
    HEX_LIGHT_BG = "F2F4F7"
    HEX_CALLOUT_BG = "EBF1F5"
    HEX_BORDER = "D3D3D3"

    # Set Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Helper Functions
    def set_run_font(run, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def set_p_spacing(p, space_before=0, space_after=6, line_spacing=1.15):
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(p, space_before=12, space_after=6)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=24, color=COLOR_NAVY, bold=True)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(p, space_before=0, space_after=24)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=13, color=COLOR_MUTED, italic=True)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=18, space_after=8)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=16, color=COLOR_NAVY, bold=True)
        # Add a light bottom border or underline visual
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=14, space_after=6)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=13, color=COLOR_SECONDARY, bold=True)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=10, space_after=4)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=11.5, color=COLOR_NAVY, bold=True)
        return p

    def add_body_p(text="", bold_prefix=""):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=0, space_after=6)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=True)
        if text:
            r_txt = p.add_run(text)
            set_run_font(r_txt, font_name="Calibri", size_pt=11, color=COLOR_DARK)
        return p

    def add_bullet(text="", bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        set_p_spacing(p, space_before=0, space_after=4)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=True)
        if text:
            r_txt = p.add_run(text)
            set_run_font(r_txt, font_name="Calibri", size_pt=11, color=COLOR_DARK)
        return p

    def add_callout(text, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)

        # Set background shading and left border
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_CALLOUT_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="36" w:space="0" w:color="{HEX_PRIMARY}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        cell._tc.get_or_add_tcPr().append(tcBorders)

        p = cell.paragraphs[0]
        set_p_spacing(p, space_before=6, space_after=6)
        if title:
            r_t = p.add_run(f"{title}\n")
            set_run_font(r_t, font_name="Calibri", size_pt=11, color=COLOR_NAVY, bold=True)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=10.5, color=COLOR_DARK, italic=True)
        doc.add_paragraph() # spacing

    def add_code_block(code_str):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_LIGHT_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            </w:tcBorders>
        ''')
        cell._tc.get_or_add_tcPr().append(tcBorders)

        p = cell.paragraphs[0]
        set_p_spacing(p, space_before=4, space_after=4, line_spacing=1.0)
        r = p.add_run(code_str)
        set_run_font(r, font_name="Consolas", size_pt=9.5, color=RGBColor(40, 40, 40))
        doc.add_paragraph()

    # --- DOCUMENT GENERATION ---

    # Document Header
    add_title("TÀI LIỆU PHÂN TÍCH, KỊCH BẢN THUYẾT TRÌNH & BỘ PHẢN BIỆN\nCHỨC NĂNG LỊCH SỬ ĐƠN HÀNG (ORDER HISTORY)")
    add_subtitle("Dự án Hệ thống Thương mại Điện tử Linh kiện Máy tính WINNOTech\nKiến trúc: Node.js / Express / MongoDB (Data Snapshot Pattern & State Machine)")

    # Metadata Box
    add_callout(
        "• Hệ quản trị CSDL: MongoDB Atlas / Mongoose ORM\n"
        "• Kiến trúc dữ liệu: Mô hình Lai (Reference + Embedded Data Snapshot Pattern)\n"
        "• Quản lý trạng thái: Máy trạng thái hữu hạn 5 bước chuẩn hóa (Canonical 5-Step Flow)\n"
        "• Phạm vi áp dụng: Bảo vệ Đồ án, Thuyết trình Dự án, Phản biện Hội đồng & Tài liệu Kỹ thuật",
        "📌 TỔNG QUAN TÀI LIỆU"
    )

    # ---------------------------------------------------------
    # PHẦN 1
    # ---------------------------------------------------------
    add_h1("PHẦN 1: KỊCH BẢN THUYẾT TRÌNH CHI TIẾT (PRESENTATION PITCH SCRIPT)")
    add_body_p("Kịch bản thuyết trình được thiết kế nhằm giúp người trình bày thể hiện sự chuyên nghiệp, mạch lạc, làm nổi bật cả giá trị nghiệp vụ (Business Value) lẫn sự vượt trội về kỹ thuật (Technical Excellence).")

    add_h2("1. Mở đầu & Đặt vấn đề (Ấn tượng & Trực diện - 30 giây)")
    add_callout(
        "\"Kính thưa Hội đồng và Qúy vị!\n"
        "Trong bất kỳ hệ thống Thương mại Điện tử chuyên nghiệp nào, trải nghiệm sau bán hàng (Post-purchase Experience) đóng vai trò quyết định đến 80% tỷ lệ quay lại của khách hàng. Chức năng Lịch sử đơn hàng (Order History) tại WINNOTech không chỉ đơn thuần là nơi hiển thị danh sách mua sắm, mà được thiết kế như một trung tâm quản trị minh bạch toàn bộ vòng đời đơn hàng, giúp người dùng theo dõi sát sao tiến trình theo thời gian thực, đồng thời tối ưu hóa công tác vận hành cho Quản trị viên.\"",
        "🎤 Lời thoại Mở đầu"
    )

    add_h2("2. Trình bày Tính năng & Giá trị Nổi bật (1.5 phút)")
    add_body_p("Dưới đây là 4 điểm sáng cốt lõi của chức năng Lịch sử đơn hàng WINNOTech:")

    add_bullet(" Hệ thống phân loại đơn hàng minh bạch qua 5 trạng thái chuẩn hóa: Chờ xác nhận ➔ Đang chuẩn bị hàng ➔ Đang giao hàng ➔ Đã giao hàng ➔ Hoàn thành, cùng luồng Đã hủy riêng biệt. Khách hàng có thể lọc đơn hàng tức thì chỉ bằng 1 cú nhấp chuột.", "1. Phân loại 5 bước Canonical Flow: ")

    add_bullet(" Mỗi khi đơn hàng thay đổi trạng thái (do Admin cập nhật hoặc Hệ thống tự động), một bản ghi vết chứa mốc thời gian, người thực hiện (Admin/Hệ thống) và ghi chú cụ thể sẽ được nhúng trực tiếp vào đơn hàng. Điều này giúp loại bỏ hoàn toàn rủi ro tranh chấp.", "2. Vết lịch sử trạng thái minh bạch (Audit Trail - statusHistory): ")

    add_bullet(" Khách hàng có thể hủy đơn linh hoạt khi đơn đang ở trạng thái Chờ xác nhận hoặc Chuẩn bị hàng. Ngay khi hủy, hệ thống tự động hoàn lại tồn kho ($inc stock_quantity) theo thời gian thực và ghi nhận yêu cầu hoàn tiền nếu đã thanh toán VNPay.", "3. Hủy đơn & Hoàn kho tự động (Auto Stock Recovery): ")

    add_bullet(" Hệ thống tự động chuyển trạng thái đơn hàng sang Hoàn thành (completed) khi đơn đã giao thành công và đã thanh toán đủ, hoặc tự động quét chốt đơn sau 10 ngày. Cơ chế này giúp chốt doanh thu chuẩn xác và kích hoạt quyền đánh giá sản phẩm.", "4. Bộ máy Tự động hóa Thông minh (Auto-Completion Engine): ")

    add_h2("3. Kết luận Lời thuyết trình (15 giây)")
    add_callout(
        "\"Tất cả các cơ chế này được xây dựng trên nền tảng cơ sở dữ liệu MongoDB tối ưu hóa theo mô hình Data Snapshot Pattern, giúp dữ liệu lịch sử tài chính bền vững và tuyệt đối không bị sai lệch theo thời gian dù sản phẩm gốc có bị chỉnh sửa hay xóa bỏ. Sau đây em xin trình bày chi tiết về kiến trúc lưu trữ và nguyên lý hoạt động bên dưới.\"",
        "🎤 Lời thoại Kết luận"
    )

    # ---------------------------------------------------------
    # PHẦN 2
    # ---------------------------------------------------------
    add_h1("PHẦN 2: CÁCH THỨC HOẠT ĐỘNG & BẢN DỮ LIỆU LƯU TRONG MONGO DB")
    add_body_p("Hệ thống WINNOTech lưu trữ dữ liệu Lịch sử đơn hàng bằng mô hình lai (Hybrid Pattern) kết hợp giữa **Reference (Tham chiếu)** và **Data Snapshot Pattern (Chụp ảnh dữ liệu tại thời điểm giao dịch)** với 2 Collections chính trong MongoDB:")

    add_h2("1. Chi tiết 2 Collections trong MongoDB")
    add_h3("A. Collection 'orders' (Lưu thông tin tổng quan, địa chỉ & tiến trình)")
    add_body_p("Được định nghĩa tại model models/Order.js. Collection này chịu trách nhiệm lưu trữ thông tin chung của đơn hàng, snapshot địa chỉ người nhận và mảng lịch sử tiến trình embedded:")

    add_code_block(
'''// Mongo Document trong Collection: orders
{
  "_id": ObjectId("65f1a2b3c4d5e6f7a8b9c0d1"),
  "user_id": ObjectId("65f000000000000000000001"), // Ref đến User
  "code": "ORD-1725400000000",                     // Mã đơn hàng duy nhất
  "status": "preparing",                           // Trạng thái hiện tại
  
  // --- SNAPSHOT THÔNG TIN NGƯỜI NHẬN (CỐ ĐỊNH) ---
  "Name": "Nguyễn Văn A",
  "Phone": "0987654321",
  "Adress": "123 Đường ABC, Phường 1, Q3, TP.HCM",
  
  // --- THÔNG TIN TÀI CHÍNH ---
  "total_amount": 25500000,                        // Tổng tiền sau giảm giá
  "payment_method": ObjectId("65f000000000000000000002"),
  "payment_status": "unpaid",                      // unpaid | paid | refund_pending | canceled
  "voucher_code": "SUMMER2026",
  "voucher_value": 500000,
  
  // --- MẢNG LỊCH SỬ TIẾN TRÌNH (STATUS HISTORY EMBEDDED) ---
  "statusHistory": [
    {
      "status": "pending",
      "note": "Tạo đơn hàng thành công",
      "changedBy": "Hệ thống",
      "changedAt": ISODate("2026-09-04T00:00:00Z")
    },
    {
      "status": "preparing",
      "note": "Xác nhận kho đủ hàng",
      "changedBy": "Admin Nguyễn Văn B",
      "changedAt": ISODate("2026-09-04T01:30:00Z")
    }
  ],
  "admin_notes": [],
  "createdAt": ISODate("2026-09-04T00:00:00Z"),
  "updatedAt": ISODate("2026-09-04T01:30:00Z")
}'''
    )

    add_h3("B. Collection 'orderitems' (Lưu danh sách sản phẩm & Snapshot giá)")
    add_body_p("Được định nghĩa tại OrderItemSchema trong models/Order.js. Lưu từng mặt hàng thuộc đơn hàng:")

    add_code_block(
'''// Mongo Document trong Collection: orderitems
{
  "_id": ObjectId("65f1a2b3c4d5e6f7a8b9c0d2"),
  "order_id": ObjectId("65f1a2b3c4d5e6f7a8b9c0d1"), // Foreign key ref đến orders
  "variants_id": ObjectId("65f999999999999999999999"), // Ref đến ProductVariant
  
  // --- SNAPSHOT GIÁ TẠI THỜI ĐIỂM MUA ---
  "Quantity": 1,
  "price": 26000000  // Giá sản phẩm tại thời điểm mua (ưu tiên sale_price)
}'''
    )

    add_h2("2. Nguyên lý Chụp ảnh Dữ liệu (Data Snapshot Pattern)")
    add_body_p("Đây là quy chuẩn cốt lõi trong phát triển ứng dụng Thương mại điện tử chuyên nghiệp:")
    add_bullet("Khi khách hàng đặt hàng thành công (API POST /orders), hệ thống truy vấn giá thực tế của linh kiện từ ProductVariant, sau đó lưu trực tiếp giá trị này vào trường orderitems.price.", "Snapshot Giá tiền (price): ")
    add_bullet("Thông tin Name, Phone, Adress được ghi thẳng vào Document Order tại thời điểm chốt đơn.", "Snapshot Người nhận: ")
    add_bullet("Cho dù sau này Admin có thay đổi giá sản phẩm từ 26 triệu xuống 20 triệu, hoặc người dùng cập nhật địa chỉ nhà mới trong Profile, hóa đơn trong Lịch sử đơn hàng vẫn giữ nguyên giá trị 26 triệu và địa chỉ cũ. Điều này bảo vệ tính toàn vẹn tài chính và pháp lý cho hóa đơn.", "Ý nghĩa nghiệp vụ: ")

    add_h2("3. Máy trạng thái (Order State Machine Flow)")
    add_body_p("Vòng đời đơn hàng tuân theo quy trình 5 bước canonical nghiêm ngặt:")
    add_body_p("pending (Chờ xác nhận) ➔ preparing (Chuẩn bị hàng) ➔ shipping (Đang giao) ➔ delivered (Đã giao) ➔ completed (Hoàn thành)")
    add_body_p("Chỉ đơn ở trạng thái pending hoặc preparing mới được chuyển sang cancelled (Đã hủy). Khi bị hủy, hệ thống tự động chạy toán tử kho:")
    add_code_block("ProductVariant.findByIdAndUpdate(variants_id, { $inc: { stock_quantity: Quantity } })")

    add_body_p("Ngoài ra, hệ thống tự động chuẩn hóa các trạng thái legacy từ dữ liệu cũ (như handed_over, shipped, delivering ➔ shipping; done ➔ completed) bằng hàm normalizeStatus() giúp giao diện không bị gián đoạn.")

    # ---------------------------------------------------------
    # PHẦN 3
    # ---------------------------------------------------------
    add_h1("PHẦN 3: BỘ CÂU HỎI & TRẢ LỜI PHẢN BIỆN (Q&A DEFENSE GUIDE)")
    add_body_p("Bộ câu hỏi phản biện chuyên sâu giúp sinh viên/nhà phát triển bảo vệ thành công dự án trước các câu hỏi xoáy từ Hội đồng giám khảo:")

    # Q1
    add_h2("Câu 1: \"Nếu Admin sửa giá sản phẩm hoặc xóa sản phẩm trong Database, lịch sử đơn hàng cũ có bị thay đổi giá hoặc lỗi hiển thị không?\"")
    add_callout(
        "\"Thưa Hội đồng, hoàn toàn KHÔNG bị ảnh hưởng ạ.\n\n"
        "Hệ thống của chúng em áp dụng thiết kế Data Snapshot Pattern trong MongoDB:\n"
        "1. Giá tiền (price), tên người nhận, số điện thoại, địa chỉ và giá trị voucher được chụp lại tại thời điểm mua và lưu cố định vào Document Order & OrderItem.\n"
        "2. Khi xem Lịch sử đơn hàng, hệ thống lấy giá tiền từ trường price trong orderitems chứ không truy vấn lại giá từ ProductVariant hiện tại.\n"
        "3. Ngay cả khi sản phẩm bị xóa khỏi database, thông tin hóa đơn và các con số tài chính vẫn được bảo toàn nguyên vẹn 100%.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q2
    add_h2("Câu 2: \"Tại sao em lại tách thành 2 Collection 'orders' và 'orderitems' trong MongoDB mà không nhúng (embed) mảng items trực tiếp vào bên trong 1 Document Order?\"")
    add_callout(
        "\"Dạ, đây là quyết định thiết kế kiến trúc dựa trên 3 lý do kỹ thuật:\n\n"
        "1. Tối ưu tốc độ tải danh sách (Listing Query): Việc tách orderitems giúp document Order giữ dung lượng nhỏ gọn. Khi người dùng tải trang danh sách Lịch sử đơn hàng, hệ thống query cực nhanh mà không bị phình dung lượng bởi mảng sub-documents.\n"
        "2. Phục vụ chức năng Đánh giá & Thống kê: Mỗi OrderItem có một _id độc lập, giúp bảng Đánh giá (Review) liên kết trực tiếp tới từng sản phẩm trong đơn qua id_oderitems.\n"
        "3. Tách biệt trách nhiệm (Separation of Concerns): Đơn hàng quản lý tiến trình giao nhận và tài chính tổng, còn orderitems quản lý chi tiết hàng hóa và giá snapshot.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q3
    add_h2("Câu 3: \"Khi người dùng hoặc Admin ấn Hủy đơn hàng, hệ thống xử lý những nghiệp vụ ngầm (Background Logic) nào?\"")
    add_callout(
        "\"Dạ, khi kích hoạt API Hủy đơn (PUT /orders/:orderId/cancel), hệ thống thực thi 4 bước nghiệp vụ liên hoàn:\n\n"
        "1. Kiểm tra điều kiện (Guard Check): Chỉ cho phép hủy khi đơn ở trạng thái pending hoặc preparing. Đơn đã giao hàng (shipping) không thể hủy tự do.\n"
        "2. Hoàn tồn kho (Stock Refund): Server tự động hoàn lại số lượng sản phẩm vào kho bằng toán tử bulk write $inc: { stock_quantity: +item.Quantity }.\n"
        "3. Xử lý trạng thái thanh toán (Payment Handling): Nếu đơn đã thanh toán VNPay (paid), payment_status chuyển thành refund_pending để Admin xử lý trả tiền; nếu là COD chuyển thành canceled.\n"
        "4. Ghi vết Audit Trail: Thêm bản ghi mới vào mảng statusHistory ghi rõ lý do hủy, người hủy và thời gian hủy.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q4
    add_h2("Câu 4: \"Tại sao em lại áp dụng Máy trạng thái (State Machine) với quy trình nghiêm ngặt mà không cho phép Admin chuyển trạng thái tùy ý?\"")
    add_callout(
        "\"Dạ, việc áp dụng Máy trạng thái (ORDER_TRANSITIONS) nhằm đảm bảo tính hợp lệ của quy trình nghiệp vụ Thương mại Điện tử:\n\n"
        "1. Ngăn chặn lỗi vận hành do con người: Admin không thể nhảy cóc từ Chờ xác nhận lên Đã giao hàng mà bỏ qua bước Chuẩn bị hàng & Giao hàng.\n"
        "2. Bảo vệ luồng tài chính: Đơn hàng chỉ được chuyển sang Hoàn thành (completed) khi và chỉ khi đơn đã ở trạng thái Đã giao hàng (delivered) VÀ đã thanh toán tiền (paid).\n"
        "3. Đảm bảo dữ liệu thống kê báo cáo chính xác.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q5
    add_h2("Câu 5: \"Dữ liệu cũ trong DB đang chứa các status legacy như handed_over, shipped, done. Hệ thống xử lý thế nào để không sụp đổ hay sai lệch?\"")
    add_callout(
        "\"Dạ, hệ thống xử lý bằng giải pháp Tương thích ngược (Backward Compatibility):\n\n"
        "1. Trong OrderSchema, giữ lại các tên trạng thái cũ trong enum validator để Mongoose không báo lỗi khi đọc dữ liệu cũ.\n"
        "2. Viết hàm chuẩn hóa normalizeStatus() trên Server: Tự động gom handed_over, shipped, delivering thành shipping; và done thành completed.\n"
        "3. Khi khách hàng bấm lọc theo tab trên giao diện, MongoDB query sử dụng toán tử $in: { status: { $in: ['shipping', 'handed_over', 'shipped', 'delivering'] } }, đảm bảo không bỏ sót bất kỳ đơn hàng nào.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q6
    add_h2("Câu 6: \"Cơ chế Tự động chuyển đơn từ delivered sang completed sau 10 ngày (hoặc khi có Đánh giá) giải quyết bài toán nghiệp vụ gì?\"")
    add_callout(
        "\"Dạ, cơ chế Auto-Completion Engine giải quyết 3 bài toán nghiệp vụ quan trọng:\n\n"
        "1. Chốt doanh thu cho cửa hàng: Tránh tình trạng đơn hàng bị treo ở trạng thái Đã giao hàng mãi mãi do người mua quên bấm 'Đã nhận được hàng'.\n"
        "2. Kích hoạt tính năng Đánh giá sản phẩm: Ngay khi khách hoàn tất đánh giá các sản phẩm trong đơn, hệ thống lập tức chốt đơn sang completed.\n"
        "3. Cập nhật số lượng đã bán (Product Sales): Khi đơn chuyển sang completed, server tự động gọi hàm incrementProductSalesForOrder() để cộng dồn số lượng bán cho linh kiện, giúp thống kê sản phẩm bán chạy chính xác.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # ---------------------------------------------------------
    # PHẦN 4
    # ---------------------------------------------------------
    add_h1("PHẦN 4: THAM CHIẾU FILE CODE TRONG DỰ ÁN WINNOTECH")
    add_body_p("Dưới đây là danh sách vị trí code triển khai thực tế trong codebase dự án:")

    # Table of files
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    hdr_cells = tbl.rows[0].cells
    headers = ["Tên File Code", "Đoạn Code / Lines", "Chức năng Nghiệp vụ"]
    widths = [Inches(2.0), Inches(1.5), Inches(3.0)]

    for idx, text in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        p = hdr_cells[idx].paragraphs[0]
        set_p_spacing(p, space_before=4, space_after=4)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=10.5, color=RGBColor(255, 255, 255), bold=True)
        # Background shading
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_PRIMARY}"/>')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shd)

    data = [
        ("models/Order.js", "Lines 4-60", "Định nghĩa OrderSchema, OrderItemSchema, statusHistory & status enum"),
        ("server.js", "Lines 2787-2946", "API POST /orders (Tạo đơn, Snapshot giá, trừ kho bulkWrite & áp voucher)"),
        ("server.js", "Lines 2951-3070", "API GET /orders (Lấy lịch sử đơn hàng, lọc status, normalize legacy status)"),
        ("server.js", "Lines 3075-3149", "API GET /orders/:orderId (Chi tiết đơn hàng, populate variants & attributes)"),
        ("server.js", "Lines 3154-3214", "API PUT /orders/:orderId/cancel (Hủy đơn hàng, hoàn stock_quantity vào kho)"),
        ("server.js", "Lines 5251-5395", "API PUT /admin/orders/:id/status (Cập nhật status, máy trạng thái & statusHistory)"),
        ("frontend/src/pages/Profile.jsx", "Frontend UI", "Giao diện Lịch sử đơn hàng, các Tab lọc trạng thái & xem chi tiết")
    ]

    for row_idx, (f_name, f_lines, f_desc) in enumerate(data):
        row_cells = tbl.add_row().cells
        bg_color = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for idx, text in enumerate([f_name, f_lines, f_desc]):
            row_cells[idx].width = widths[idx]
            p = row_cells[idx].paragraphs[0]
            set_p_spacing(p, space_before=4, space_after=4)
            r = p.add_run(text)
            is_bold = (idx == 0)
            set_run_font(r, font_name="Calibri", size_pt=10, color=COLOR_DARK, bold=is_bold)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            row_cells[idx]._tc.get_or_add_tcPr().append(shd)
            # cell borders
            tcBorders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                </w:tcBorders>
            ''')
            row_cells[idx]._tc.get_or_add_tcPr().append(tcBorders)

    # Output file
    output_filename = "Phan_Tich_Va_Phan_Bien_Lich_Su_Don_Hang_WINNOTech.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}")

if __name__ == "__main__":
    create_document()
