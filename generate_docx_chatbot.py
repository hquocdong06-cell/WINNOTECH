import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_chatbot_document():
    doc = Document()

    # Define Theme Colors
    COLOR_NAVY = RGBColor(31, 78, 120)     # #1F4E78 - Primary Heading
    COLOR_SECONDARY = RGBColor(46, 117, 182) # #2E75B6 - Secondary Heading
    COLOR_DARK = RGBColor(38, 38, 38)       # #262626 - Text
    COLOR_MUTED = RGBColor(89, 89, 89)     # #595959 - Subtitle / Muted

    HEX_PRIMARY = "1F4E78"
    HEX_LIGHT_BG = "F2F4F7"
    HEX_CALLOUT_BG = "EBF1F5"
    HEX_BORDER = "D3D3D3"

    # Set Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Helper Functions
    def set_run_font(run, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def set_p_spacing(p, space_before=0, space_after=6, line_spacing=1.15):
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(p, space_before=12, space_after=6)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=24, color=COLOR_NAVY, bold=True)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(p, space_before=0, space_after=24)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=13, color=COLOR_MUTED, italic=True)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=18, space_after=8)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=16, color=COLOR_NAVY, bold=True)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=14, space_after=6)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=13, color=COLOR_SECONDARY, bold=True)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=10, space_after=4)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=11.5, color=COLOR_NAVY, bold=True)
        return p

    def add_body_p(text="", bold_prefix=""):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=0, space_after=6)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=True)
        if text:
            r_txt = p.add_run(text)
            set_run_font(r_txt, font_name="Calibri", size_pt=11, color=COLOR_DARK)
        return p

    def add_bullet(text="", bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        set_p_spacing(p, space_before=0, space_after=4)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=True)
        if text:
            r_txt = p.add_run(text)
            set_run_font(r_txt, font_name="Calibri", size_pt=11, color=COLOR_DARK)
        return p

    def add_callout(text, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)

        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_CALLOUT_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="36" w:space="0" w:color="{HEX_PRIMARY}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        cell._tc.get_or_add_tcPr().append(tcBorders)

        p = cell.paragraphs[0]
        set_p_spacing(p, space_before=6, space_after=6)
        if title:
            r_t = p.add_run(f"{title}\n")
            set_run_font(r_t, font_name="Calibri", size_pt=11, color=COLOR_NAVY, bold=True)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=10.5, color=COLOR_DARK, italic=True)
        doc.add_paragraph()

    def add_code_block(code_str):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)

        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_LIGHT_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            </w:tcBorders>
        ''')
        cell._tc.get_or_add_tcPr().append(tcBorders)

        p = cell.paragraphs[0]
        set_p_spacing(p, space_before=4, space_after=4, line_spacing=1.0)
        r = p.add_run(code_str)
        set_run_font(r, font_name="Consolas", size_pt=9.5, color=RGBColor(40, 40, 40))
        doc.add_paragraph()

    # --- DOCUMENT GENERATION ---

    # Document Header
    add_title("TÀI LIỆU PHÂN TÍCH, KỊCH BẢN THUYẾT TRÌNH & BỘ PHẢN BIỆN\nCHỨC NĂNG 11: CHATBOT AI TƯ VẤN LINH KIỆN & CẤU HÌNH PC")
    add_subtitle("Dự án Hệ thống Thương mại Điện tử Linh kiện Máy tính WINNOTech\nKiến trúc: Google Gemini AI, RAG Database Context, Multi-Model Fallback Hierarchy & Dynamic Backup Reply")

    # Metadata Box
    add_callout(
        "• AI Core Engine: Google Gemini API (@google/generative-ai)\n"
        "• Mô hình RAG: Retrieval-Augmented Generation (Truy vấn MongoDB Kho hàng tạo ngữ cảnh sản phẩm thực tế)\n"
        "• Cơ chế Tin cậy: Chuỗi Fallback Đa mô hình (Multi-Model Hierarchy) + Bộ phản hồi dự phòng động (Dynamic Fallback Reply)\n"
        "• Nguyên lý Kiến trúc: Chatbot Stateless (Không lưu chat vào DB, Frontend quản lý history theo phiên)\n"
        "• Phạm vi áp dụng: Bảo vệ Đồ án, Thuyết trình Dự án, Phản biện Hội đồng & Tài liệu Kỹ thuật",
        "📌 TỔNG QUAN CHỨC NĂNG CHATBOT AI"
    )

    # ---------------------------------------------------------
    # CHƯƠNG 1
    # ---------------------------------------------------------
    add_h1("CHƯƠNG 1: KỊCH BẢN THUYẾT TRÌNH CHI TIẾT (AI CHATBOT PITCH SCRIPT)")
    add_body_p("Kịch bản thuyết trình được thiết kế nhằm giúp người trình bày làm nổi bật giá trị đột phá của Trợ lý AI trong việc cá nhân hóa trải nghiệm khách hàng và tự động hóa tư vấn bán hàng 24/7.")

    add_h2("1.1 Mở đầu & Đặt vấn đề (Ấn tượng & Trực diện - 30 giây)")
    add_callout(
        "\"Kính thưa Hội đồng và Qúy vị!\n"
        "Trong ngành bán lẻ Linh kiện máy tính và PC Gaming, việc chọn mua đúng loại CPU phù hợp với Socket Mainboard, đúng công suất Nguồn PSU hay vừa vặn với túi tiền là rào cản lớn nhất đối với khách hàng. Chức năng Chatbot AI tại WINNOTech được xây dựng dựa trên công nghệ Google Gemini AI, đóng vai trò như một Chuyên gia tư vấn phần cứng 24/7. AI không chỉ trả lời câu hỏi thông thường mà còn suy luận thông minh, tính toán ngân sách và trực tiếp đề xuất các sản phẩm thực tế đang có sẵn tại showroom WINNOTech.\"",
        "🎤 Lời thoại Mở đầu"
    )

    add_h2("1.2 Trình bày 4 Điểm sáng Cốt lõi & Giá trị Nghiệp vụ (1.5 phút)")
    add_body_p("Dưới đây là 4 trụ cột công nghệ vượt trội của Chức năng Chatbot AI WINNOTech:")

    add_bullet(" Backend truyền System Instruction chuyên sâu đóng vai tư vấn viên phần cứng. AI tự phân tích logic yêu cầu của khách hàng, suy luận mục đích sử dụng (gaming, đồ họa 3D, lập trình, stream, văn phòng), giải thích khả năng tương thích kỹ thuật (Socket LGA1700/AM5, DDR4/DDR5, PCIe 4.0/5.0, công suất PSU) và trình bày bằng Markdown đẹp mắt với icon sinh động.", "1. Tích hợp Google Gemini AI & System Instruction Chuyên sâu: ")

    add_bullet(" Hệ thống tích hợp kỹ thuật RAG (Retrieval-Augmented Generation): Backend tự động trích xuất ngân sách (triệu, k, ngàn) và từ khóa linh kiện từ câu hỏi của người dùng, truy vấn các sản phẩm thực tế từ Database MongoDB, sau đó đính kèm dữ liệu kho hàng vào ngữ cảnh cho AI. Điều này giúp AI ưu tiên giới thiệu sản phẩm thật kèm giá chính xác của WINNOTech thay vì trả lời chung chung.", "2. Mô hình RAG - Đưa Dữ liệu Kho hàng Thực tế vào Ngữ cảnh AI: ")

    add_bullet(" Để đảm bảo chatbot luôn hoạt động 100% không bị gián đoạn, hệ thống thiết lập chuỗi thử nghiệm liên tục qua 4 model Gemini mới nhất (gemini-3.5-flash-lite ➔ gemini-3.6-flash ➔ gemini-3.1-pro-preview ➔ gemini-2.5-flash). Nếu API Gemini hoặc mạng gặp sự cố, hệ thống lập tức kích hoạt Bộ trả lời dự phòng động (Dynamic Fallback Reply) dựa trên MongoDB và các chính sách bảo hành/giao hàng.", "3. Cơ chế Thử nhiều Model & Bộ Trả lời Dự phòng Động (Fault-Tolerant Engine): ")

    add_bullet(" Chatbot hoạt động hoàn toàn Stateless — không lưu lịch sử chat vào Database nhằm giải phóng tài nguyên lưu trữ và bảo vệ quyền riêng tư của khách hàng. Frontend tự quản lý mảng tin nhắn và chuẩn hóa history truyền về cho Backend trong mỗi lượt chat. Giao diện dạng Floating Widget tích hợp bóng thoại chào mừng, bộ 3 nút mạng xã hội (Messenger, Zalo, Chatbot AI) và thanh câu hỏi gợi ý nhanh.", "4. Kiến trúc Stateless Tối ưu & Giao diện UI/UX Hiện đại: ")

    add_h2("1.3 Kết luận Lời thuyết trình (15 giây)")
    add_callout(
        "\"Tất cả các cơ chế này tạo nên một hệ thống tư vấn AI vừa thông minh, vừa thực tế với dữ liệu kho hàng, lại có độ tin cậy tuyệt đối nhờ cơ chế dự phòng đa lớp. Sau đây em xin trình bày chi tiết về luồng xử lý dữ liệu và bộ câu hỏi phản biện bên dưới.\"",
        "🎤 Lời thoại Kết luận"
    )

    # ---------------------------------------------------------
    # CHƯƠNG 2
    # ---------------------------------------------------------
    add_h1("CHƯƠNG 2: CÁCH THỨC HOẠT ĐỘNG & KIẾN TRÚC KỸ THUẬT BACKEND / FRONTEND")
    add_body_p("Kiến trúc kỹ thuật của Chatbot AI WINNOTech bao gồm 5 bước xử lý End-to-End được thiết kế tối ưu hiệu năng:")

    add_h2("2.1 Luồng Xử lý Dữ liệu End-to-End (Data Flow Pipeline)")
    add_bullet("Client gửi request POST /api/chatbot/chat kèm payload { message, history }.", "Step 1 (User Query): ")
    add_bullet("Backend chạy hàm parseBudgetFromMessage() để nhận diện số tiền (ví dụ: '15tr' ➔ 15,000,000đ; '500k' ➔ 500,000đ) và tìm kiếm sản phẩm trong MongoDB qua hàm findRelevantProducts().", "Step 2 (Budget & Product RAG Search): ")
    add_bullet("Chuẩn hóa mảng history gửi từ client: Chuyển đổi role thành 'user'/'model', trích xuất text từ parts, lọc bỏ tin nhắn trống và đảm bảo tin nhắn đầu tiên trong mảng luôn có role là 'user'.", "Step 3 (History Normalization): ")
    add_bullet("Đính kèm danh sách sản phẩm thực tế trong kho vào System Instruction của Gemini. Khởi tạo GoogleGenerativeAI và thực hiện vòng lặp thử lần lượt các model trong candidateModels.", "Step 4 (Gemini AI Execution with System Instruction): ")
    add_bullet("Nếu một model trả về đáp án ➔ Trả về JSON { success: true, reply, source: 'gemini' }. Nếu tất cả model AI bị lỗi ➔ Gọi buildDynamicFallbackReply() và trả về { success: true, reply, source: 'dynamic_fallback' }.", "Step 5 (Fallback Response Guarantee): ")

    add_h2("2.2 Minh họa Đoạn Code Xử lý Cốt lõi tại Backend (AI_chatbot.js)")
    add_code_block(
'''// 1. Hàm Trích xuất Ngân sách từ câu hỏi tự nhiên
function parseBudgetFromMessage(msg) {
  const text = msg.toLowerCase().replace(/,/g, "").replace(/\\./g, "");
  const millionMatch = text.match(/(\\d+(?:\\.\\d+)?)\\s*(?:triệu|trieu|tr|m\\b)/i);
  if (millionMatch) return parseFloat(millionMatch[1]) * 1000000;
  const thousandMatch = text.match(/(\\d+)\\s*(?:k|ngàn|ngan|nghìn|nghin)/i);
  if (thousandMatch) return parseInt(thousandMatch[1], 10) * 1000;
  return null;
}

// 2. Chuỗi thử nghiệm Multi-Model Fallback Hierarchy
const candidateModels = [
  "gemini-3.5-flash-lite",
  "gemini-3.6-flash",
  "gemini-3.1-pro-preview",
  "gemini-2.5-flash"
];

for (const modelName of candidateModels) {
  try {
    const model = genAI.getGenerativeModel({
      model: modelName,
      systemInstruction: `Bạn là Trợ lý AI chuyên gia tư vấn phần cứng PC & Siêu thị công nghệ WINNOTech... ${productContext}`
    });
    const chat = model.startChat({ history: validHistory });
    const result = await chat.sendMessage(trimmedMsg);
    aiReply = result.response.text();
    if (aiReply) break;
  } catch (modelErr) {
    console.warn(`Thử model ${modelName} thất bại:`, modelErr.message);
  }
}'''
    )

    add_h2("2.3 Nguyên lý Thiết kế Stateless & Lợi ích Hệ thống")
    add_body_p("Chatbot được thiết kế theo mô hình **Stateless Architecture**:")
    add_bullet("Dữ liệu hội thoại không lưu vào bất kỳ Collection nào trong MongoDB. Mỗi lượt chat là độc lập về phía Server.", "• Khái niệm Stateless: ")
    add_bullet("Tiết kiệm dung lượng bộ nhớ Database, tránh làm phình RAM/Disk IO, không tốn chi phí dọn dẹp log chat cũ.", "• Tối ưu Tài nguyên DB: ")
    add_bullet("Khách hàng vãng lai (Guest User) có thể trò chuyện tư vấn ngay lập tức mà không cần đăng nhập tài khoản.", "• Trải nghiệm Người dùng (UX): ")
    add_bullet("Tránh rủi ro rò rỉ thông tin cá nhân hoặc dữ liệu riêng tư của người dùng trong CSDL.", "• Bảo mật & Quyền riêng tư: ")

    # ---------------------------------------------------------
    # CHƯƠNG 3
    # ---------------------------------------------------------
    add_h1("CHƯƠNG 3: BỘ CÂU HỎI & TRẢ LỜI PHẢN BIỆN CHUYÊN SÂU (Q&A DEFENSE GUIDE)")
    add_body_p("Bộ câu hỏi phản biện chuyên sâu giúp sinh viên/nhà phát triển tự tin bảo vệ tính năng trước Hội đồng giám khảo:")

    # Q1
    add_h2("Câu 1: \"Tại sao Chatbot lại hoạt động theo cơ chế Stateless mà không lưu lịch sử cuộc trò chuyện vào MongoDB?\"")
    add_callout(
        "\"Thưa Hội đồng, đây là quyết định thiết kế kiến trúc dựa trên 3 lý do kỹ thuật:\n\n"
        "1. Tối ưu Tài nguyên Database: Việc lưu hàng triệu tin nhắn chat vãng lai sẽ làm phình dung lượng CSDL và gây quá tải I/O ghi đĩa không cần thiết.\n"
        "2. Hỗ trợ Khách hàng Vãng lai (Guest Users): Cho phép mọi khách hàng truy cập website đều có thể sử dụng chatbot tư vấn ngay lập tức mà không bắt buộc đăng nhập.\n"
        "3. Tách biệt Trách nhiệm (Client-Side History Management): Frontend ReactJS chịu trách nhiệm lưu giữ mảng messages trong React State phiên làm việc và truyền kèm payload history trong mỗi request, đảm bảo Backend luôn xử lý đúng ngữ cảnh mà không cần lưu trữ state.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q2
    add_h2("Câu 2: \"Hệ thống khắc phục hiện tượng AI 'chém gió' (hallucination) hoặc tư vấn các sản phẩm không có bán tại cửa hàng bằng cách nào?\"")
    add_callout(
        "\"Dạ, chúng em áp dụng mô hình RAG (Retrieval-Augmented Generation):\n\n"
        "1. Khi người dùng hỏi (ví dụ: 'tư vấn VGA dưới 10 triệu'), Backend lập tức phân tích từ khóa và tìm kiếm trong Collection Product & ProductVariant của WINNOTech.\n"
        "2. Danh sách sản phẩm thực tế kèm giá bán chính xác và tình trạng tồn kho sẽ được nhúng thẳng vào System Instruction của Gemini.\n"
        "3. Trong System Instruction, chúng em quy định nghiêm ngặt: 'Hãy ưu tiên giới thiệu các sản phẩm thật đang có sẵn tại kho WINNOTech'. Nhờ đó, AI trả lời chuẩn xác 100% dựa trên dữ liệu thực của cửa hàng.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q3
    add_h2("Câu 3: \"Khi Google Gemini API gặp sự cố (quá tải quota, mất mạng hoặc API Key lỗi), Chatbot xử lý ra sao để không bị crash trang web?\"")
    add_callout(
        "\"Dạ, hệ thống áp dụng Cơ chế Phòng thủ 2 Lớp (Two-Layer Fallback Architecture):\n\n"
        "1. Lớp 1 - Multi-Model Fallback Hierarchy: Server thiết lập mảng candidateModels gồm 4 phiên bản Gemini mới nhất (gemini-3.5-flash-lite, gemini-3.6-flash, gemini-3.1-pro-preview, gemini-2.5-flash). Nếu model này lỗi quota, vòng lặp tự động chuyển sang model tiếp theo trong milisecond.\n"
        "2. Lớp 2 - Dynamic Fallback Reply: Nếu tất cả các model AI đều không phản hồi, server sẽ chuyển sang hàm buildDynamicFallbackReply(). Hàm này tự động trích xuất sản phẩm từ DB MongoDB hoặc đưa ra câu trả lời chuẩn về chính sách bảo hành/giao hàng của WINNOTech. Nhờ vậy, người dùng LUÔN nhận được phản hồi hữu ích mà không bao giờ bị báo lỗi hệ thống.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q4
    add_h2("Câu 4: \"Backend chuẩn hóa mảng 'history' gửi từ Frontend như thế nào trước khi truyền vào Gemini API?\"")
    add_callout(
        "\"Dạ, hàm handleChat thực hiện 3 bước chuẩn hóa history nghiêm ngặt:\n\n"
        "1. Mapping Role: Quy đổi role 'human' hoặc 'user' thành 'user', và các role khác thành 'model'.\n"
        "2. Trích xuất Nội dung (Text Extraction): Đảm bảo lấy đúng text từ item.parts[0].text hoặc item.text và loại bỏ các phần tử tin nhắn rỗng.\n"
        "3. Lọc Tin nhắn Đầu tiên: Vòng lặp while kiểm tra nếu phần tử đầu tiên của validHistory không phải là 'user' thì sẽ shift bỏ, đảm bảo luồng hội thoại bắt đầu đúng chuẩn quy định của Google Gemini SDK.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q5
    add_h2("Câu 5: \"Làm sao Chatbot hiểu được ngân sách của người dùng khi người dùng gõ câu hỏi dạng tự nhiên (ví dụ: 'tầm 15tr', 'dưới 2 triệu', '500k')?\"")
    add_callout(
        "\"Dạ, chúng em xây dựng thuật toán trích xuất biểu thức chính quy (Regex Pattern Matching) trong hàm parseBudgetFromMessage():\n\n"
        "1. Xử lý hàng Triệu: Regex /(\\d+(?:\\.\\d+)?)\\s*(?:triệu|trieu|tr|m)/i trích xuất số và nhân với 1,000,000 (vd: '15tr' ➔ 15,000,000).\n"
        "2. Xử lý hàng Nghìn: Regex /(\\d+)\\s*(?:k|ngàn|ngan|nghìn|nghin)/i trích xuất số và nhân với 1,000 (vd: '500k' ➔ 500,000).\n"
        "3. Xử lý Số thô: Trích xuất các chuỗi số nguyên từ 6-10 chữ số. Số tiền trích xuất sẽ được dùng để query lọc sản phẩm có giá price <= budget * 1.1 trong MongoDB.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q6
    add_h2("Câu 6: \"Giao diện Frontend Chatbot UI được thiết kế tối ưu trải nghiệm người dùng (UX) như thế nào?\"")
    add_callout(
        "\"Dạ, component AIChatbot.jsx được thiết kế chỉn chu với 4 điểm nhấn UX:\n\n"
        "1. Vertical Floating Stack: Cụm 3 nút nổi đồng nhất kích thước (52px x 52px) gồm Messenger, Zalo và Chatbot AI đặt ở góc phải dưới màn hình.\n"
        "2. Greeting Pill: Bóng thoại 'Bạn cần hỗ trợ gì?' tự động hiển thị gợi ý kích thích người dùng bấm tương tác.\n"
        "3. Quick Suggestions: Thanh cuộn ngang chứa các câu hỏi mẫu ('Tư vấn PC 15 triệu chơi game', 'Chọn PSU phù hợp RTX 3060') giúp gửi câu hỏi nhanh chỉ bằng 1 cú nhấp chuột.\n"
        "4. Auto Scroll & Markdown Styling: Tự động cuộn xuống tin nhắn mới nhất và định dạng văn bản rõ ràng.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # ---------------------------------------------------------
    # CHƯƠNG 4
    # ---------------------------------------------------------
    add_h1("CHƯƠNG 4: BẢNG THAM CHIẾU FILE CODE THỰC TẾ TRONG DỰ ÁN WINNOTECH")
    add_body_p("Tổng hợp toàn bộ các vị trí file code triển khai thực tế trong codebase dự án:")

    # Table of files
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    hdr_cells = tbl.rows[0].cells
    headers = ["Tên File Code", "Đoạn Code / Lines", "Chức năng Nghiệp vụ"]
    widths = [Inches(2.2), Inches(1.5), Inches(2.8)]

    for idx, text in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        p = hdr_cells[idx].paragraphs[0]
        set_p_spacing(p, space_before=4, space_after=4)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=10.5, color=RGBColor(255, 255, 255), bold=True)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_PRIMARY}"/>')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shd)

    master_data = [
        ("routers/AI_chatbot.js", "Lines 19-41", "Hàm parseBudgetFromMessage (Trích xuất ngân sách VNĐ từ tin nhắn tự nhiên)"),
        ("routers/AI_chatbot.js", "Lines 46-138", "Hàm findRelevantProducts (RAG Query sản phẩm phù hợp từ MongoDB)"),
        ("routers/AI_chatbot.js", "Lines 143-181", "Hàm buildDynamicFallbackReply (Bộ trả lời dự phòng động từ DB & chính sách)"),
        ("routers/AI_chatbot.js", "Lines 186-312", "Hàm handleChat (Chuẩn hóa history, System Instruction, Multi-Model Fallback)"),
        ("routers/AI_chatbot.js", "Lines 314-315", "Khai báo Route POST /chat và POST /ask"),
        ("server.js", "Lines 8318-8329", "Mount Router /api/chatbot, /chatbot, /api/chat vào ứng dụng Express"),
        ("frontend/src/components/AIChatbot.jsx", "Lines 36-76", "Hàm handleSend (Gửi câu hỏi, quản lý history state & loading status)"),
        ("frontend/src/components/AIChatbot.jsx", "Lines 89-350", "Giao diện Popup Chat Window (Header, Messages List, Quick Suggestions, Input)"),
        ("frontend/src/components/AIChatbot.jsx", "Lines 352-578", "Cụm Vertical Floating Stack (Nút Messenger, Zalo, Chatbot AI & Greeting Pill)"),
        ("frontend/src/services/apiService.js", "Lines 366-377", "Export chatbotAPI.chat() và chatbotAPI.ask() gọi về Backend Express")
    ]

    for row_idx, (f_name, f_lines, f_desc) in enumerate(master_data):
        row_cells = tbl.add_row().cells
        bg_color = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for idx, text in enumerate([f_name, f_lines, f_desc]):
            row_cells[idx].width = widths[idx]
            p = row_cells[idx].paragraphs[0]
            set_p_spacing(p, space_before=4, space_after=4)
            r = p.add_run(text)
            is_bold = (idx == 0)
            set_run_font(r, font_name="Calibri", size_pt=10, color=COLOR_DARK, bold=is_bold)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            row_cells[idx]._tc.get_or_add_tcPr().append(shd)
            tcBorders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                </w:tcBorders>
            ''')
            row_cells[idx]._tc.get_or_add_tcPr().append(tcBorders)

    # Output file
    output_filename = "Phan_Tich_Va_Phan_Bien_AI_Chatbot_WINNOTech.docx"
    doc.save(output_filename)
    print(f"Successfully generated chatbot document {output_filename}")

if __name__ == "__main__":
    create_chatbot_document()
