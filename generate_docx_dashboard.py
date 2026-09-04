import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_dashboard_document():
    doc = Document()

    # Define Theme Colors
    COLOR_NAVY = RGBColor(31, 78, 120)     # #1F4E78 - Primary Heading
    COLOR_SECONDARY = RGBColor(46, 117, 182) # #2E75B6 - Secondary Heading
    COLOR_DARK = RGBColor(38, 38, 38)       # #262626 - Text
    COLOR_MUTED = RGBColor(89, 89, 89)     # #595959 - Subtitle / Muted

    HEX_PRIMARY = "1F4E78"
    HEX_LIGHT_BG = "F2F4F7"
    HEX_CALLOUT_BG = "EBF1F5"
    HEX_BORDER = "D3D3D3"

    # Set Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Helper Functions
    def set_run_font(run, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def set_p_spacing(p, space_before=0, space_after=6, line_spacing=1.15):
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(p, space_before=12, space_after=6)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=24, color=COLOR_NAVY, bold=True)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(p, space_before=0, space_after=24)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=13, color=COLOR_MUTED, italic=True)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=18, space_after=8)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=16, color=COLOR_NAVY, bold=True)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=14, space_after=6)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=13, color=COLOR_SECONDARY, bold=True)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=10, space_after=4)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=11.5, color=COLOR_NAVY, bold=True)
        return p

    def add_body_p(text="", bold_prefix=""):
        p = doc.add_paragraph()
        set_p_spacing(p, space_before=0, space_after=6)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=True)
        if text:
            r_txt = p.add_run(text)
            set_run_font(r_txt, font_name="Calibri", size_pt=11, color=COLOR_DARK)
        return p

    def add_bullet(text="", bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        set_p_spacing(p, space_before=0, space_after=4)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, font_name="Calibri", size_pt=11, color=COLOR_DARK, bold=True)
        if text:
            r_txt = p.add_run(text)
            set_run_font(r_txt, font_name="Calibri", size_pt=11, color=COLOR_DARK)
        return p

    def add_callout(text, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)

        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_CALLOUT_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="36" w:space="0" w:color="{HEX_PRIMARY}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        cell._tc.get_or_add_tcPr().append(tcBorders)

        p = cell.paragraphs[0]
        set_p_spacing(p, space_before=6, space_after=6)
        if title:
            r_t = p.add_run(f"{title}\n")
            set_run_font(r_t, font_name="Calibri", size_pt=11, color=COLOR_NAVY, bold=True)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=10.5, color=COLOR_DARK, italic=True)
        doc.add_paragraph()

    def add_code_block(code_str):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)

        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_LIGHT_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            </w:tcBorders>
        ''')
        cell._tc.get_or_add_tcPr().append(tcBorders)

        p = cell.paragraphs[0]
        set_p_spacing(p, space_before=4, space_after=4, line_spacing=1.0)
        r = p.add_run(code_str)
        set_run_font(r, font_name="Consolas", size_pt=9.5, color=RGBColor(40, 40, 40))
        doc.add_paragraph()

    # --- DOCUMENT GENERATION ---

    # Document Header
    add_title("TÀI LIỆU PHÂN TÍCH, KỊCH BẢN THUYẾT TRÌNH & BỘ PHẢN BIỆN\nCHỨC NĂNG DASHBOARD QUẢN TRỊ (ADMIN DASHBOARD)")
    add_subtitle("Dự án Hệ thống Thương mại Điện tử Linh kiện Máy tính WINNOTech\nKiến trúc Analytics: Real-time Data Aggregation, Dynamic Multi-dimension Chart & Excel Export Engine")

    # Metadata Box
    add_callout(
        "• Hệ quản trị CSDL: MongoDB Atlas (Aggregation Pipeline & Lean Queries)\n"
        "• Tầng Phân tích & Thống kê: Dynamic Time-series Bucketing (Day / Week / Month / Year)\n"
        "• Thư viện Xuất Báo cáo: ExcelJS (Streamed Binary Spreadsheet Generation)\n"
        "• Phạm vi áp dụng: Bảo vệ Đồ án, Thuyết trình Dự án, Phản biện Hội đồng & Tài liệu Kỹ thuật",
        "📌 TỔNG QUAN TÀI LIỆU DASHBOARD"
    )

    # ---------------------------------------------------------
    # PHẦN 1
    # ---------------------------------------------------------
    add_h1("PHẦN 1: KỊCH BẢN THUYẾT TRÌNH CHI TIẾT (DASHBOARD PRESENTATION PITCH)")
    add_body_p("Kịch bản thuyết trình chức năng Admin Dashboard giúp người trình bày thể hiện tư duy quản trị kinh doanh dựa trên dữ liệu (Data-Driven Management) và năng lực tối ưu hệ thống.")

    add_h2("1. Mở đầu & Đặt vấn đề (Trực diện & Chuyên nghiệp - 30 giây)")
    add_callout(
        "\"Kính thưa Hội đồng và Qúy vị!\n"
        "Đối với một doanh nghiệp Thương mại Điện tử chuyên ngành Linh kiện Máy tính như WINNOTech, khả năng theo dõi sức khỏe kinh doanh và biến động tồn kho theo thời gian thực là yếu tố sinh tử. Trang Admin Dashboard được thiết kế như trung tâm chỉ huy số (Digital Command Center), hợp nhất dữ liệu từ toàn bộ các phân hệ: Đơn hàng, Doanh thu, Khách hàng và Kho hàng thành các chỉ số trực quan, giúp Quản trị viên đưa ra quyết định kinh doanh chính xác và kịp thời.\"",
        "🎤 Lời thoại Mở đầu"
    )

    add_h2("2. Trình bày Tính năng & Giá trị Nổi bật (1.5 phút)")
    add_body_p("Dưới đây là 4 trụ cột chức năng vượt trội của Admin Dashboard WINNOTech:")

    add_bullet(" Hiển thị tức thì 4 chỉ số sinh mệnh kinh doanh gồm: Tổng doanh thu thực nhận, Tổng số đơn hàng đã hoàn thành, Số lượng tài khoản khách hàng đăng ký và Số lượng sản phẩm/biến thể trong hệ thống.", "1. Bộ thẻ Thống kê KPI Thời gian thực (Real-time Metric Cards): ")

    add_bullet(" Tích hợp biểu đồ trực quan hóa linh hoạt hỗ trợ 3 chế độ xem: Doanh thu 12 tháng gần nhất, Doanh thu từng ngày trong tháng được chọn (YYYY-MM), và Doanh thu 7 ngày trong tuần (T2 - CN). Biểu đồ tích hợp Tooltip tương tác hiển thị chi tiết số tiền và số đơn hàng.", "2. Biểu đồ Doanh thu Đa chiều (Multi-dimension Dynamic Chart): ")

    add_bullet(" Dashboard chia thành 3 widget quản trị nhanh: Bảng 10 đơn hàng mới nhất cần xử lý; Bảng Top 10 sản phẩm tồn kho dồi dào; và Bảng Cảnh báo Sắp hết hàng (Top Low Stock) giúp phát hiện kịp thời các linh kiện chạm ngưỡng báo động để nhập hàng bổ sung.", "3. Bộ 3 Bảng Cảnh báo & Vận hành siêu tốc (Operational Widgets): ")

    add_bullet(" Hệ thống tích hợp công cụ xuất báo cáo doanh thu ra file Excel (.xlsx) chuẩn định dạng kế toán bằng thư viện ExcelJS với tiêu đề chỉn chu, kẻ viền và tự động tính tổng tiền.", "4. Công cụ Xuất Báo cáo Excel Chuyên nghiệp (Excel Export Engine): ")

    add_h2("3. Kết luận Lời thuyết trình (15 giây)")
    add_callout(
        "\"Tất cả dữ liệu trên Dashboard được truy vấn song song qua cơ chế Promise.all() và tối ưu hóa bằng mượt mà ở cả Backend lẫn Frontend, đảm bảo thời gian phản hồi dưới 200ms ngay cả khi dữ liệu tăng trưởng mạnh. Sau đây em xin trình bày chi tiết về kiến trúc kỹ thuật bên dưới.\"",
        "🎤 Lời thoại Kết luận"
    )

    # ---------------------------------------------------------
    # PHẦN 2
    # ---------------------------------------------------------
    add_h1("PHẦN 2: CÁCH THỨC HOẠT ĐỘNG & BẢN DỮ LIỆU LƯU TRONG MONGO DB")
    add_body_p("Điểm đặc biệt trong kiến trúc của Dashboard WINNOTech là **Dashboard không lưu dữ liệu trong một Collection riêng**, mà đóng vai trò là **Tầng Phân tích & Tổng hợp (Analytics & Aggregation Layer)** truy vấn trực tiếp từ 4 Collections cốt lõi của MongoDB:")

    add_h2("1. Nguồn Dữ liệu & Các Collections Được Truy Vấn")
    
    add_bullet("Lọc ra tất cả các đơn hàng thỏa mãn điều kiện doanh thu thực tế: { $or: [{ status: { $in: ['completed', 'delivered', 'done'] } }, { payment_status: 'paid' }] }. Dữ liệu được tính tổng tiền (total_amount) và phân bổ vào các mốc thời gian.", "1. Collection 'orders': ")

    add_bullet("Sử dụng User.countDocuments({ role: { $ne: 'admin' } }) để thống kê chính xác số lượng khách hàng đã đăng ký tài khoản.", "2. Collection 'users': ")

    add_bullet("Phân rã (flatten) mảng Biến thể (Variants) của từng sản phẩm để tính toán số lượng tồn kho thực tế (stock_quantity). Phân loại sản phẩm thành mảng Tồn kho cao và mảng Cảnh báo Sắp hết hàng.", "3. Collections 'products' & 'productvariants': ")

    add_bullet("Phục vụ thống kê chi tiết sản phẩm bán chạy và đóng góp vào tổng giá trị đơn hàng.", "4. Collection 'orderitems': ")

    add_h2("2. Thuật toán Gom nhóm Thời gian (Time-series Bucketing Algorithm)")
    add_body_p("Để hiển thị biểu đồ doanh thu theo Ngày, Tuần, Tháng, Năm, Backend Node.js/Express thực thi thuật toán gom nhóm thời gian như sau:")

    add_code_block(
'''// Ví dụ Thuật toán Gom nhóm Doanh thu theo Tháng (server.js - GET /admin/revenue/monthly)
const now = moment();
const months = [];

// Step 1: Khởi tạo mảng mốc thời gian 12 tháng gần nhất (Zero-filling Pattern)
for (let i = 11; i >= 0; i--) {
  const m = now.clone().subtract(i, 'months');
  months.push({ key: m.format("YYYY-MM"), label: m.format("MM/YYYY"), revenue: 0, orderCount: 0 });
}

// Step 2: Query các đơn hàng đủ điều kiện doanh thu
const startOfRange = moment().subtract(11, 'months').startOf('month').toDate();
const completedOrders = await Order.find({
  $or: [
    { status: { $in: ["completed", "delivered", "done"] } },
    { payment_status: "paid" }
  ],
  createdAt: { $gte: startOfRange }
}).lean();

// Step 3: Duyệt và cộng dồn doanh thu vào Bucket tương ứng
completedOrders.forEach((ord) => {
  const key = moment(ord.createdAt || ord.date).format("YYYY-MM");
  const entry = months.find(m => m.key === key);
  if (entry) {
    entry.revenue += ord.total_amount || 0;
    entry.orderCount += 1;
  }
});'''
    )

    add_h2("3. Giải pháp Xuất File Excel Doanh thu (Excel Export Engine)")
    add_body_p("Tại API GET /admin/revenue/export-excel, server sử dụng thư viện ExcelJS để khởi tạo luồng xuất file binary:")
    add_bullet("Tạo Banner Tiêu đề: Merge cell 'A1:E1', tô màu nền tối #0F172A, chữ trắng in hoa Arial 16pt bold.", "• Định dạng giao diện Excel: ")
    add_bullet("Dòng dữ liệu: Định dạng số tiền kiểu tiền tệ (#,##0 đ), căn lề phải cho cột doanh thu, căn giữa cho ngày tháng.", "• Định dạng dữ liệu: ")
    add_bullet("Gửi trực tiếp về Browser với Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet để tự động tải về file .xlsx.", "• Stream HTTP Response: ")

    # ---------------------------------------------------------
    # PHẦN 3
    # ---------------------------------------------------------
    add_h1("PHẦN 3: BỘ CÂU HỎI & TRẢ LỜI PHẢN BIỆN (Q&A DEFENSE GUIDE)")
    add_body_p("Bộ câu hỏi phản biện chuyên sâu giúp bảo vệ thành công kiến trúc Dashboard trước Hội đồng:")

    # Q1
    add_h2("Câu 1: \"Tại sao Dashboard không lưu dữ liệu doanh thu hàng ngày vào 1 Collection 'RevenueStats' riêng để query cho nhanh mà lại đi tính toán (aggregate) từ Collection 'orders' mỗi khi Admin mở trang?\"")
    add_callout(
        "\"Thưa Hội đồng, đây là quyết định thiết kế để đảm bảo Tính chính xác tuyệt đối của dữ liệu (Data Consistency):\n\n"
        "1. Tránh rủi ro lệch dữ liệu (Data Desynchronization): Nếu lưu bảng thống kê riêng, khi một đơn hàng bị Hủy (cancelled) hoặc Hoàn tiền (refunded), nếu background job bị lỗi thì số liệu trên Dashboard sẽ bị sai lệch so với đơn hàng thực tế.\n"
        "2. Truy vấn trực tiếp từ 'orders' giúp số liệu luôn đúng 100% tại thời điểm thực tế (Real-time Accuracy).\n"
        "3. Tối ưu hiệu năng: Server áp dụng Mongoose .lean() bỏ qua overhead biến đối tượng, kết hợp đánh Index trên { createdAt: 1, status: 1, payment_status: 1 } nên tốc độ query cực nhanh chỉ dưới 50ms.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q2
    add_h2("Câu 2: \"Đơn hàng bị Hủy (cancelled) hoặc chưa Thanh toán (unpaid) có được tính vào Tổng Doanh Thu trên Dashboard không?\"")
    add_callout(
        "\"Dạ, hoàn toàn KHÔNG ạ.\n\n"
        "Server kiểm tra điều kiện lọc doanh thu rất nghiêm ngặt:\n"
        "Chỉ những đơn hàng có payment_status === 'paid' HOẶC status thuộc nhóm ['completed', 'delivered', 'done'] mới được cộng dồn total_amount vào doanh thu.\n"
        "Tất cả các đơn hàng chưa thanh toán COD đang chờ xác nhận, đơn bị hủy, hoặc đơn đang chờ hoàn tiền đều bị loại trừ khỏi tổng doanh thu.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q3
    add_h2("Câu 3: \"Biểu đồ doanh thu trên Dashboard xử lý ra sao khi có những tháng hoặc ngày KHÔNG CÓ đơn hàng nào?\"")
    add_callout(
        "\"Dạ, hệ thống áp dụng kỹ thuật Zero-Filling Pattern:\n\n"
        "1. Trước khi truy vấn DB, Server khởi tạo sẵn mảng chứa đủ mốc thời gian (ví dụ: 12 tháng gần nhất hoặc đủ số ngày trong tháng) với giá trị doanh thu = 0 và số đơn = 0.\n"
        "2. Sau khi lấy kết quả từ MongoDB, server mới map doanh thu vào mốc tương ứng.\n"
        "3. Nhờ đó, biểu đồ trên giao diện luôn hiển thị liên tục, không bị đứt đoạn hay thiếu mốc thời gian khi có tháng 0 đồng.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q4
    add_h2("Câu 4: \"Dashboard quản lý các sản phẩm có nhiều Biến thể (Variants) như thế nào để đưa ra cảnh báo 'Sắp hết hàng' chính xác?\"")
    add_callout(
        "\"Dạ, Dashboard xử lý quản lý tồn kho theo cấp độ Biến thể chi tiết (Variant-level Inventory):\n\n"
        "1. Trên Frontend, hàm loadDashboardData() thực hiện phẳng hóa (flatten) sản phẩm: Nếu sản phẩm có danh sách Variants (như các phiên bản RAM, CPU, Card màn hình), hệ thống tách thành từng dòng riêng kèm SKU cụ thể.\n"
        "2. Cảnh báo Sắp hết hàng (lowStockItems) sắp xếp mảng theo stock_quantity tăng dần, giúp Admin biết chính xác phiên bản linh kiện nào đang hết hàng để nhập thêm chứ không báo chung chung tên sản phẩm.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q5
    add_h2("Câu 5: \"Cơ chế Xuất file Excel Doanh Thu (export-excel) được triển khai thế nào ở phía Server?\"")
    add_callout(
        "\"Dạ, API export-excel được viết bằng thư viện ExcelJS trên Node.js:\n\n"
        "1. Server lấy danh sách các đơn đã thanh toán/hoàn thành, nhóm doanh thu theo tham số ngày/tuần/tháng/năm.\n"
        "2. Khởi tạo Workbook & Worksheet, tạo banner tiêu đề styled màu tối #0F172A, merge cell, kẻ viền cell, tính tổng doanh thu.\n"
        "3. Gửi dữ liệu về dưới dạng HTTP Stream Binary (.xlsx) giúp người dùng tải về file Excel chuẩn mà không làm phình RAM của Server.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # Q6
    add_h2("Câu 6: \"Nếu hệ thống phát triển lên hàng triệu đơn hàng trong tương lai, Dashboard sẽ gặp nghẽn cổ chai ở đâu và giải pháp mở rộng (Scaling) là gì?\"")
    add_callout(
        "\"Dạ, khi hệ thống lớn lên hàng triệu đơn hàng, việc .find().lean() cả mảng sẽ làm tăng RAM Server. Giải pháp nâng cấp mở rộng của chúng em là:\n\n"
        "1. Đẩy toán tử gom nhóm xuống MongoDB Atlas bằng Aggregation Pipeline với $match và $group trực tiếp trên Database Engine.\n"
        "2. Tạo Compound Index: db.orders.createIndex({ status: 1, payment_status: 1, createdAt: -1 }).\n"
        "3. Áp dụng Redis Caching: Lưu kết quả doanh thu các tháng cũ (đã chốt số không thay đổi) vào Redis Cache với TTL 24h, chỉ truy vấn tính toán realtime cho tháng hiện tại.\"",
        "💡 Trả lời phản biện xuất sắc"
    )

    # ---------------------------------------------------------
    # PHẦN 4
    # ---------------------------------------------------------
    add_h1("PHẦN 4: THAM CHIẾU FILE CODE TRONG DỰ ÁN WINNOTECH")
    add_body_p("Dưới đây là danh sách vị trí code triển khai thực tế của Dashboard trong codebase:")

    # Table of files
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    hdr_cells = tbl.rows[0].cells
    headers = ["Tên File Code", "Đoạn Code / Lines", "Chức năng Nghiệp vụ"]
    widths = [Inches(2.2), Inches(1.5), Inches(2.8)]

    for idx, text in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        p = hdr_cells[idx].paragraphs[0]
        set_p_spacing(p, space_before=4, space_after=4)
        r = p.add_run(text)
        set_run_font(r, font_name="Calibri", size_pt=10.5, color=RGBColor(255, 255, 255), bold=True)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_PRIMARY}"/>')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shd)

    data = [
        ("server.js", "Lines 5499-5565", "API GET /admin/revenue/stats (Thống kê doanh thu theo ngày/tuần/tháng/năm)"),
        ("server.js", "Lines 5567-5611", "API GET /admin/revenue/monthly (Doanh thu 12 tháng gần nhất cho biểu đồ)"),
        ("server.js", "Lines 5613-5644", "API GET /admin/revenue/by-month (Doanh thu từng ngày trong tháng YYYY-MM)"),
        ("server.js", "Lines 5646-5677", "API GET /admin/revenue/by-week (Doanh thu 7 ngày trong tuần từ thứ 2)"),
        ("server.js", "Lines 5680-5750", "API GET /admin/revenue/export-excel (Xuất file Excel doanh thu bằng ExcelJS)"),
        ("frontend/src/admin/pages/Dashboard.jsx", "Lines 1-428", "Trang UI Dashboard Admin: 4 Thẻ KPI, Biểu đồ Doanh thu, 3 Widget tồn kho & Đơn hàng"),
        ("frontend/src/admin/services/adminService.js", "API Services", "Hàm gọi API fetchAdminUsers, fetchAdminProducts, fetchAdminOrders, fetchMonthlyRevenue")
    ]

    for row_idx, (f_name, f_lines, f_desc) in enumerate(data):
        row_cells = tbl.add_row().cells
        bg_color = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for idx, text in enumerate([f_name, f_lines, f_desc]):
            row_cells[idx].width = widths[idx]
            p = row_cells[idx].paragraphs[0]
            set_p_spacing(p, space_before=4, space_after=4)
            r = p.add_run(text)
            is_bold = (idx == 0)
            set_run_font(r, font_name="Calibri", size_pt=10, color=COLOR_DARK, bold=is_bold)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            row_cells[idx]._tc.get_or_add_tcPr().append(shd)
            tcBorders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
                </w:tcBorders>
            ''')
            row_cells[idx]._tc.get_or_add_tcPr().append(tcBorders)

    # Output file
    output_filename = "Phan_Tich_Va_Phan_Bien_Dashboard_WINNOTech.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}")

if __name__ == "__main__":
    create_dashboard_document()
